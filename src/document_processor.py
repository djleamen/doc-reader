"""
Document processing utilities for various file formats.
"""
import os
import magic
from typing import List, Dict, Any
from pathlib import Path
from loguru import logger

import pypdf
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

from src.config import settings


class DocumentProcessor:
    """Handles document loading, processing, and chunking."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_document(self, file_path: str) -> str:
        """Load and extract text from various document formats."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > settings.max_document_size_mb:
            raise ValueError(f"Document too large: {file_size_mb:.2f}MB (max: {settings.max_document_size_mb}MB)")
        
        file_extension = file_path.suffix.lower().lstrip('.')
        
        if file_extension not in settings.supported_formats_list:
            raise ValueError(f"Unsupported format: {file_extension}")
        
        logger.info(f"Loading document: {file_path}")
        
        if file_extension == 'pdf':
            return self._load_pdf(file_path)
        elif file_extension == 'docx':
            return self._load_docx(file_path)
        elif file_extension in ['txt', 'md']:
            return self._load_text(file_path)
        else:
            raise ValueError(f"Handler not implemented for: {file_extension}")
    
    def _load_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
        
        return text.strip()
    
    def _load_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
    
    def _load_text(self, file_path: Path) -> str:
        """Load plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with latin-1 encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
    
    def chunk_document(self, text: str, metadata: Dict[str, Any] = None) -> List[LangChainDocument]:
        """Split document into chunks for vector storage."""
        if metadata is None:
            metadata = {}
        
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                **metadata,
                "chunk_id": i,
                "chunk_size": len(chunk)
            }
            documents.append(LangChainDocument(
                page_content=chunk,
                metadata=chunk_metadata
            ))
        
        logger.info(f"Created {len(documents)} chunks from document")
        return documents
    
    def process_document(self, file_path: str, additional_metadata: Dict[str, Any] = None) -> List[LangChainDocument]:
        """Complete document processing pipeline."""
        file_path = Path(file_path)
        
        # Load document text
        text = self.load_document(file_path)
        
        # Prepare metadata
        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "file_type": file_path.suffix.lower().lstrip('.'),
            "file_size": file_path.stat().st_size,
            "word_count": len(text.split())
        }
        
        if additional_metadata:
            metadata.update(additional_metadata)
        
        # Chunk the document
        documents = self.chunk_document(text, metadata)
        
        return documents
